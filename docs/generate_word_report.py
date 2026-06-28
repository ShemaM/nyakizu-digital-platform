from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from pathlib import Path
import os
import textwrap
import fnmatch
import io

out_path = Path('docs/ui-db-code-design-report.docx')
doc = Document()

def title_page(title, course, author, date):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(title + '\n')
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph()
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p2.add_run(course + '\n')
    p2.add_run(author + '\n')
    p2.add_run(date + '\n')
    # Insert a page break
    doc.add_page_break()

def add_section_heading(text):
    doc.add_heading(text, level=1)

def add_subheading(text):
    doc.add_heading(text, level=2)

def add_paragraph(text):
    para = doc.add_paragraph(text)
    para.runs[0].font.size = Pt(11)

def explain_file(path: Path) -> str:
    """Return a short heuristic explanation for a file's purpose based on keywords."""
    try:
        text = path.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        return 'Could not read file content.'
    snippet = textwrap.shorten(text.replace('\n', ' '), width=300, placeholder='...')
    lower = text.lower()
    explanation = 'Auto-explanation: ' \
        + ('React/Next component or frontend module.' if 'export default' in lower or 'react' in lower else '')
    if 'class ' in text and 'models' in path.parts:
        explanation = 'Defines database models for the application domain.'
    if 'def create' in lower and 'orders' in str(path).lower():
        explanation = 'Implements order creation business logic and stock updates.'
    if 'def' in lower and 'serializer' in str(path).lower():
        explanation = 'Serializers for API input validation and (some) business logic.'
    if 'service worker' in lower or 'self.addEventListener' in lower or 'sw.js' in str(path).lower():
        explanation = 'Service worker implementing caching strategies and offline fallback.'
    if 'next' in str(path).lower() or path.suffix in ('.ts','.tsx'):
        if 'api.' in str(path).lower() or 'lib/api' in str(path).lower():
            explanation = 'Frontend API client: centralizes network calls and domain types.'
    if explanation == 'Auto-explanation: ':
        explanation = 'General source file. See repository for full implementation.'
    return f"{explanation}\n\nFile snippet: {snippet}"

def add_repo_explanations(root_dirs):
    add_subheading('Repository code explanations (per-file)')
    files_found = []
    for root in root_dirs:
        for dirpath, dirnames, filenames in os.walk(root):
            for pattern in ('*.py','*.tsx','*.ts','*.js'):
                for fname in fnmatch.filter(filenames, pattern):
                    p = Path(dirpath) / fname
                    rel = p.relative_to(Path.cwd())
                    files_found.append(rel)
                    doc.add_heading(str(rel), level=3)
                    expl = explain_file(p)
                    add_paragraph(expl)
    if not files_found:
        add_paragraph('No source files found under the specified roots.')

def add_screenshot_placeholders():
    add_section_heading('Screenshots')
    add_paragraph('This report will include screenshots demonstrating key UI flows. Insert the final screenshots into the `docs/screenshots/` folder and re-run the generator to embed them. For now the report includes textual placeholders to indicate where images should go.')
    add_paragraph('Expected screenshot placeholders (replace with real images named similarly):')
    add_paragraph('- [Insert screenshot: buyer-home - include Windows taskbar time in the frame]')
    add_paragraph('- [Insert screenshot: buyer-product-list - include Windows taskbar time in the frame]')
    add_paragraph('- [Insert screenshot: seller-dashboard - include Windows taskbar time in the frame]')
    add_paragraph('- [Insert screenshot: seller-order-detail - include Windows taskbar time in the frame]')
    add_paragraph('- [Insert screenshot: login-screen - include Windows taskbar time in the frame]')
    add_paragraph('When you add real screenshots, name them using this pattern: `<role>-<screen>-YYYYMMDD-HHMMSS.png` and re-run the script.')

# Title page (APA-like simple)
title_page('SWE 3090XA User Interface, DATABASE and Code Design Report', 'SWE 3090XA — Software Engineering Project I', 'Author: Nzabakamira Shema Manasse', 'June 2026')

# Abstract
doc.add_heading('Abstract', level=1)
add_paragraph('This report provides an academic evaluation of the Nyakizu Digital Market project focusing on user interface design, database schema and backend code design. The evaluation maps implemented artifacts to the project requirements and identifies gaps, risks, and recommendations. Findings are grounded in repository inspection and design documentation. Keywords: Nyakizu, PWA, offline-first, ledger, seller-buyer relationship.')
doc.add_page_break()

# Follow the exact section numbering and headings requested by the user
add_section_heading('1. Introduction')
add_subheading('Purpose of the report')
add_paragraph('This report evaluates the current implementation of the Nyakizu Digital Market project with respect to its user interface, database, and code design. It documents observed behavior, identifies deviations from the project concept, and provides prioritized recommendations for addressing functional and non-functional gaps.')
add_subheading('Scope')
add_paragraph('The scope encompasses UI evaluation and code design analysis. The UI evaluation inspects layout, consistency, accessibility, responsiveness, and usability as implemented in the frontend. The code design analysis inspects architecture, modularity, scalability, maintainability, and security aspects of the backend and API layer.')
add_subheading('Target audience')
add_paragraph('The intended audience includes project developers, UI/UX designers, and academic stakeholders responsible for validating technical robustness and usability.')

add_section_heading('2. User Interface Analysis')
add_subheading('Layout')
add_paragraph('The application uses a role-based shell to arrange navigation elements and primary page regions. The desktop view exposes a left sidebar for navigation and account controls; mobile uses a bottom navigation bar. Primary actions (suppliers, shopping lists, orders, ledger) are reachable within one or two taps from the shell. Forms (registration, login, order creation) appear as central cards with clear action buttons. See `frontend/src/components/AppShell.tsx` and `frontend/src/app/login/page.tsx` for implementation details.')
add_subheading('Consistency')
add_paragraph('The system applies a consistent dark theme with brand accents. Utility-first classes provide consistent spacing and typography. Shared UI primitives under `frontend/src/components/ui/` are reused to maintain visual consistency across roles and screens. Color usage and status badges are documented in `docs/brand-identity.md` and applied across the layout and product cards.')
add_subheading('Accessibility')
add_paragraph('The implementation demonstrates basic accessibility patterns: labeled inputs, ARIA-friendly controls (e.g., password visibility toggle), and semantic heading structure. However, a formal WCAG assessment is not present; contrast ratios and screen-reader flow require verification. Recommended tests include color-contrast measurement, keyboard navigation checks, and screen-reader walkthroughs for critical flows (registration, ordering, ledger).')
add_subheading('Responsiveness')
add_paragraph('The UI is mobile-first and responsive: viewport metadata and adaptive navigation enable usage on phones, tablets, and desktops. Major views collapse or change navigation placement by breakpoint, satisfying the project requirement to support budget smartphones.')
add_subheading('Usability')
add_paragraph('The login and onboarding flows are straightforward with clear error surfaces mapped from backend responses. The shell design reduces feature discovery friction via role-targeted navigation. Offline indicators and status labels are defined in the design docs; the service worker provides an initial offline fallback but the full offline sync UX remains incomplete.')

add_section_heading('3. Code Design Analysis')
add_subheading('Architecture')
add_paragraph('The backend follows a layered monolith architecture implemented with Django and Django REST Framework. The frontend uses Next.js App Router with React and TypeScript. The separation of frontend and backend via REST endpoints supports a clear API boundary but remains a single deployable application per layer (not microservices).')
add_subheading('Modularity')
add_paragraph('The codebase is organized into domain apps (`accounts`, `products`, `orders`) which provide a good separation of concerns. Frontend concerns (layout, API client, components) are modularized. Business logic is currently mixed into serializers (notably `orders/serializers.py`), reducing reusability and making unit testing harder. A service-layer refactor is recommended for complex operations.')
add_subheading('Scalability')
add_paragraph('Data access patterns use Django ORM with joins on related tables (e.g., SellerOrderList using items__product__seller). These patterns scale with proper indexing but may require pagination tuning, DB indices on foreign keys, and sharding or caching for high order throughput. Offline sync design will also influence scalability due to idempotency and reconciliation requirements.')
add_subheading('Maintainability')
add_paragraph('Code is readable and uses typed contracts on the frontend. However, critical domain logic embedded in serializers complicates maintenance. Tests and formal documentation are limited in the repository; adding unit tests for service routines and API integration tests would improve maintainability.')
add_subheading('Security')
add_paragraph('Authentication uses session cookies with CSRF exemptions for development convenience. Production should enable strict CSRF protection and secure cookies (HTTPS). Input validation exists in serializers, but authorization checks (e.g., enforcing buyer-seller approval at order creation) must be hardened. Sensitive fields (seller stock_quantity) are present and must be filtered from buyer-facing serializers to protect privacy.')

add_section_heading('4. Integration of UI and Code')
add_subheading('Front-end and back-end interaction')
add_paragraph('The frontend relies on a centralized API client (`frontend/src/lib/api.ts`) that defines endpoints and domain types. Endpoints align with backend views (`backend/accounts/views.py`, `backend/orders/views.py`). The UI expects server-enforced invariants (seller approval, order locking) that are only partially implemented on the backend. Tight coupling exists where the UI assumes certain server behaviors; this should be documented and tested.')
add_subheading('API usage and data flow')
add_paragraph('API flows include registration/login, seller listing, product browsing, order creation, seller order retrieval, and ledger queries. The request/response patterns are JSON-based with centralized error parsing. Order creation currently snapshots product prices and decrements stock within the serializer; this has implications for eventual consistency when offline drafts are synced.')
add_subheading('Error handling between layers')
add_paragraph('Backend validation errors are returned as structured JSON and surfaced via `ApiError` on the frontend. This approach supports consistent user feedback but requires stable error schemas and client-side handling for offline replays and idempotent operations.')

add_section_heading('5. Recommendations')
add_subheading('UI recommendations')
add_paragraph('1. Complete the offline UX: implement an IndexedDB-based sync queue, display clear status badges (Draft, Pending Sync, Synced), and use idempotency keys for retryable operations.')
add_paragraph('2. Accessibility: perform full WCAG 2.1 AA audits, ensure color contrast, keyboard operability, and screen-reader semantics for critical flows.')
add_paragraph('3. Usability: add inline confirmations and undo affordances for destructive actions (e.g., cancel order), and ensure error messages are contextual and actionable.')
add_subheading('Code recommendations')
add_paragraph('1. Refactor domain logic into service modules (e.g., `orders/services.py`) to centralize order creation, locking, and ledger updates.')
add_paragraph('2. Add append-only ledger and payment transaction models; enforce corrections as new entries to preserve auditability.')
add_paragraph('3. Enforce buyer-seller relationship approval at the API level and add unit/integration tests for this behavior.')
add_paragraph('4. Harden authentication and CSRF settings for production; store sensitive API behavior in server-side checks rather than client assumptions.')

add_section_heading('6. Conclusion')
add_subheading('Summary of findings')
add_paragraph('The project demonstrates a coherent UI design, role-respecting navigation, and a functional backend skeleton. Critical concept features (append-only ledger, robust offline sync, enforced buyer-seller gating, and order locking) are planned and partially scaffolded but require focused backend and client work to meet the project requirements.')
add_subheading('Final thoughts on usability and technical robustness')
add_paragraph('With prioritized implementation of offline sync, ledger models, and service-layer refactoring, the existing code and UI foundation can meet the project goals. The recommended changes are actionable within the current architecture and should be validated through tests and a staged deployment plan.')

# Save document
out_path.parent.mkdir(parents=True, exist_ok=True)
tmp_path = out_path.with_suffix('.tmp.docx')
try:
    doc.save(tmp_path)
    try:
        # Attempt atomic replace
        os.replace(tmp_path, out_path)
        print(out_path)
    except PermissionError:
        print(f'Could not overwrite {out_path}. It may be open. Saved temporary file at: {tmp_path}')
except Exception as e:
    print('Failed to save document:', e)